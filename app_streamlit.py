import os
import shutil
from pathlib import Path
import sys
from io import BytesIO
import streamlit as st

try:
    from docx import Document  # python-docx
except Exception:
    Document = None

# Ensure we can import from src/ as a package
sys.path.insert(0, str(Path(__file__).parent))

st.set_page_config(page_title="Medical Diagnosis AI", page_icon="🏥", layout="wide")


@st.cache_resource(show_spinner=False)
def get_system():
    from src.main import MedicalDiagnosisSystem
    return MedicalDiagnosisSystem()


def ensure_dirs():
    base = Path(__file__).parent
    (base / "data" / "vector_store").mkdir(parents=True, exist_ok=True)
    (base / "data" / "medical_textbooks").mkdir(parents=True, exist_ok=True)


def load_vector_store_from_hf(repo_id: str, subfolder=None, repo_type=None) -> str:
    """Download a prebuilt FAISS vector store from a Hugging Face repo into data/vector_store/medical_knowledge.

    Expects the repo (dataset or space or model) to contain the folder 'medical_knowledge' with index.faiss and index.pkl.
    """
    ensure_dirs()
    try:
        from huggingface_hub import snapshot_download, login
    except Exception as e:
        raise RuntimeError(f"huggingface_hub not available: {e}")

    # Token (optional for public repos)
    token = (
        os.getenv("HUGGING_FACE_HUB_TOKEN")
        or os.getenv("HUGGINGFACEHUB_API_TOKEN")
        or os.getenv("HF_TOKEN")
        or os.getenv("HF_API_HEY")  # custom secret name
    )
    if token:
        try:
            login(token=token)
        except Exception:
            pass

    # repo_type can be 'dataset', 'model', or 'space'; None lets the hub infer but is less reliable
    local_dir = snapshot_download(repo_id=repo_id, repo_type=repo_type, allow_patterns=None)
    src_root = Path(local_dir)
    if subfolder:
        src_root = src_root / subfolder

    src = src_root / "medical_knowledge"
    if not src.exists():
        # Try direct files in subfolder
        if (src_root / "index.faiss").exists() and (src_root / "index.pkl").exists():
            src = src_root
        else:
            raise FileNotFoundError("Could not find 'medical_knowledge' folder or FAISS files (index.faiss/index.pkl) in the specified repo/subfolder.")

    dst = Path(__file__).parent / "data" / "vector_store" / "medical_knowledge"
    if dst.exists():
        shutil.rmtree(dst)
    shutil.copytree(src, dst)
    return str(dst)


st.title("🏥 Medical Diagnosis AI (Streamlit)")
st.markdown("This Space uses free Hugging Face models and a medical knowledge base.")

# Optional: auto-sync vector store from Hugging Face on startup (avoids storing big files in GitHub)
try:
    VEC_DIR = Path(__file__).parent / "data" / "vector_store" / "medical_knowledge"
    if not VEC_DIR.exists():
        repo_env = os.getenv("HF_VECTOR_STORE_REPO")  # e.g. "username/medical_kb_repo"
        repo_type_env = os.getenv("HF_REPO_TYPE") or "dataset"  # dataset | model | space
        subfolder_env = os.getenv("HF_SUBFOLDER")  # optional
        if repo_env:
            with st.spinner("Syncing vector store from Hugging Face (one-time)..."):
                path = load_vector_store_from_hf(repo_id=repo_env, subfolder=subfolder_env, repo_type=repo_type_env)
            # Recreate system to pick up the new index
            get_system.clear()
            _ = get_system()
            st.success(f"Vector store ready: {path}")
except Exception as _e:
    st.info("Vector store not auto-synced. You can sync manually from the sidebar if needed.")

with st.sidebar:
    st.header("⚙️ Tools")

    # Reset session moved to sidebar
    if st.button("Reset Session", use_container_width=True):
        st.session_state.diag_state = None
        st.session_state.last_question_id = None
        st.session_state.chat_history = []
        st.rerun()

    st.divider()
    with st.expander("Advanced: Sync vector store from Hugging Face"):
        st.caption("Use this if you want to download a prebuilt FAISS index from your HF repo. If your database is already in /data/vector_store, you can ignore this.")
        repo = st.text_input("HF repo id", help="Format: namespace/name, e.g., username/medical_kb_repo")
        repo_type = st.selectbox("Repo type", ["dataset", "model", "space"], index=0, help="The type of Hugging Face repository. Choose 'dataset' if you stored the index in a dataset repo.")
        sub = st.text_input("Optional subfolder", help="Folder under the repo that contains 'medical_knowledge' or the index files.")
        if st.button("Download & Load", use_container_width=True, disabled=not repo.strip()):
            try:
                with st.spinner("Downloading vector store from Hugging Face..."):
                    path = load_vector_store_from_hf(repo_id=repo.strip(), subfolder=sub.strip() or None, repo_type=repo_type)
                st.success(f"Vector store synced to: {path}")
                st.info("Recreating system to pick up the new index...")
                get_system.clear()
                _sys = get_system()
                st.success("Vector store ready.")
            except Exception as e:
                st.error(f"Failed to load vector store: {e}")

    st.divider()
    st.subheader("📊 Status")
    if st.button("Show Status"):
        try:
            status = get_system().get_system_status()
            st.json(status)
        except Exception as e:
            st.error(f"Status error: {e}")

    # Inline controls are rendered next to the chat when needed


"""Top-of-page Patient Information form (part of the main text flow)."""
st.markdown("---")
st.markdown("#### 👤 Patient Information (optional)")
pi = st.session_state.get("patient_info") or {}
col_pi1, col_pi2 = st.columns(2)
with col_pi1:
    age = st.number_input("Age", min_value=0, max_value=130, value=int(pi.get("age", 30)), key="pi_age_top")
    gender = st.selectbox(
        "Gender",
        ["Not specified", "Male", "Female", "Other"],
        index=["Not specified", "Male", "Female", "Other"].index(pi.get("gender", "Not specified")),
        key="pi_gender_top",
    )
with col_pi2:
    med_hist = st.text_area("Medical history", value=pi.get("medical_history", ""), height=80, key="pi_history_top")
    meds = st.text_area("Current medications", value=pi.get("medications", ""), height=60, key="pi_meds_top")
    allergies = st.text_input("Allergies", value=pi.get("allergies", ""), key="pi_allergies_top")

col_piA, col_piB = st.columns([1, 1])
with col_piA:
    if st.button("Save Patient Info", use_container_width=True, key="pi_save_top"):
        st.session_state.patient_info = {
            "age": age,
            "gender": gender,
            "medical_history": med_hist,
            "medications": meds,
            "allergies": allergies,
        }
        st.success("Saved.")
with col_piB:
    if st.button("Clear", use_container_width=True, key="pi_clear_top"):
        st.session_state.patient_info = {}
        st.success("Cleared.")

# Patient info dict for downstream calls
patient = st.session_state.get("patient_info") or {}


# Diagnosis chat (full conversation view)
st.subheader("💬 Consultation")

# Session state init
if "diag_state" not in st.session_state:
    st.session_state.diag_state = None
if "last_question_id" not in st.session_state:
    st.session_state.last_question_id = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []  # list of {role: "user"|"assistant", content: str}
if "patient_info" not in st.session_state:
    st.session_state.patient_info = {}
if "show_patient_form" not in st.session_state:
    st.session_state.show_patient_form = False

# Render chat history
for msg in st.session_state.chat_history:
    with st.chat_message(msg.get("role", "assistant")):
        st.markdown(msg.get("content", ""))

# Helper: append a message
def append_message(role: str, content: str):
    st.session_state.chat_history.append({"role": role, "content": content})

# Ensure the current question (if any) appears once in chat
def show_current_question_if_new(state_dict: dict):
    q = (state_dict or {}).get("question", {})
    qid = q.get("id")
    qtext = q.get("text")
    if qtext and qid and qid != st.session_state.last_question_id:
        append_message("assistant", qtext)
        st.session_state.last_question_id = qid

# Chat input (works for both first symptoms and subsequent answers)
user_input = st.chat_input("Type your symptoms to start, then answer questions here…")
if user_input:
    sysobj = get_system()
    # New conversation: treat input as symptoms
    if not st.session_state.diag_state:
        append_message("user", user_input)
        with st.spinner("Analyzing symptoms…"):
            res = sysobj.workflow.start_interactive_diagnosis(user_input, patient)
        st.session_state.diag_state = res
        if isinstance(res, dict) and res.get("status") == "question_pending":
            show_current_question_if_new(res)
        elif isinstance(res, dict) and res.get("status") == "diagnosis_complete":
            final = res.get("final_diagnosis", {})
            summary_lines = []
            if final.get("primary_diagnosis"):
                summary_lines.append(f"Primary diagnosis: {final.get('primary_diagnosis')}")
            if final.get("final_summary"):
                summary_lines.append(final.get("final_summary"))
            if summary_lines:
                append_message("assistant", "\n\n".join(summary_lines))
        elif isinstance(res, dict) and res.get("status") == "error":
            append_message("assistant", f"Error: {res.get('error')}")
            st.error(f"Diagnosis failed: {res.get('error')}")
        st.rerun()
    else:
        # Ongoing Q&A: treat input as answer to the pending question
        state = st.session_state.diag_state
        status = (state or {}).get("status")
        append_message("user", user_input)
        if status == "question_pending":
            q = state.get("question", {})
            try:
                with st.spinner("Processing your answer…"):
                    next_state = sysobj.workflow.answer_question(q.get("id", ""), user_input, state.get("state"))
                st.session_state.diag_state = next_state
                if isinstance(next_state, dict) and next_state.get("status") == "question_pending":
                    show_current_question_if_new(next_state)
                elif isinstance(next_state, dict) and next_state.get("status") == "diagnosis_complete":
                    final = next_state.get("final_diagnosis", {})
                    summary_lines = []
                    if final.get("primary_diagnosis"):
                        summary_lines.append(f"Primary diagnosis: {final.get('primary_diagnosis')}")
                    if final.get("final_summary"):
                        summary_lines.append(final.get("final_summary"))
                    meds_section = next_state.get("medications", {}).get("suggestions", [])
                    meds_lines = [m.get("suggestion") for m in meds_section[:5] if m.get("suggestion")]
                    if meds_lines:
                        summary_lines.append("\n**Treatment suggestions:**\n- " + "\n- ".join(meds_lines))
                    if summary_lines:
                        append_message("assistant", "\n\n".join(summary_lines))
                st.rerun()
            except Exception as e:
                append_message("assistant", f"Error processing answer: {e}")
                st.rerun()
        else:
            # If conversation is completed, start a new one
            append_message("assistant", "Session complete. Click Reset Session to start over, or type new symptoms to begin a new session.")

# Controls
colA, colB = st.columns([1, 1])
with colA:
    st.caption("Use the chat to converse with the AI. It will ask follow-up questions.")
with colB:
    if st.button("Edit patient info", use_container_width=True):
        st.session_state.show_patient_form = True
        st.rerun()

### Patient Information (moved to top) — removed bottom duplicate form

# Inline structured controls for the current question (context-aware UI)
state_now = st.session_state.get("diag_state")
if isinstance(state_now, dict) and state_now.get("status") == "question_pending":
    q = state_now.get("question", {})
    qid = q.get("id", "q")
    qtext = q.get("text", "")
    qtype = (q.get("type") or "text").lower()
    opts = q.get("options") or []
    qmin = q.get("min") if q.get("min") is not None else 0
    qmax = q.get("max") if q.get("max") is not None else 10
    qdef = q.get("default")

    # Only show UI controls when needed (non-text types)
    if qtype in {"slider", "radio", "select", "multiselect", "number", "date"}:
        st.markdown("---")
        st.caption("Answer using a control or type in the chat above.")
        with st.container(border=True):
            value = None
            if qtype == "slider":
                default_val = int(qdef) if isinstance(qdef, (int, float)) else int((qmin + qmax) / 2)
                value = st.slider(qtext or "Select a value", min_value=int(qmin), max_value=int(qmax), value=default_val, key=f"slider_{qid}")
            elif qtype == "radio":
                options = opts or ["Yes", "No"]
                value = st.radio(qtext or "Choose one", options, key=f"radio_{qid}")
            elif qtype == "select":
                options = opts or ["Option 1", "Option 2"]
                value = st.selectbox(qtext or "Select one", options, key=f"select_{qid}")
            elif qtype == "multiselect":
                options = opts or ["Option A", "Option B"]
                value = st.multiselect(qtext or "Select any", options, default=qdef or [], key=f"multi_{qid}")
            elif qtype == "number":
                default_val = int(qdef) if isinstance(qdef, (int, float)) else int(qmin)
                value = st.number_input(qtext or "Enter a number", min_value=int(qmin), max_value=int(qmax), value=default_val, key=f"num_{qid}")
            elif qtype == "date":
                value = st.date_input(qtext or "Pick a date", key=f"date_{qid}")

            col1, col2 = st.columns([1, 2])
            with col1:
                if st.button("Submit", key=f"submit_{qid}", use_container_width=True):
                    # Convert control value to string answer
                    if isinstance(value, list):
                        answer_text = ", ".join(map(str, value))
                    else:
                        answer_text = str(value)
                    if answer_text.strip():
                        append_message("user", answer_text)
                        try:
                            with st.spinner("Processing your answer…"):
                                next_state = get_system().workflow.answer_question(q.get("id", ""), answer_text, state_now.get("state"))
                            st.session_state.diag_state = next_state
                            if isinstance(next_state, dict) and next_state.get("status") == "question_pending":
                                show_current_question_if_new(next_state)
                            elif isinstance(next_state, dict) and next_state.get("status") == "diagnosis_complete":
                                final = next_state.get("final_diagnosis", {})
                                summary_lines = []
                                if final.get("primary_diagnosis"):
                                    summary_lines.append(f"Primary diagnosis: {final.get('primary_diagnosis')}")
                                if final.get("final_summary"):
                                    summary_lines.append(final.get("final_summary"))
                                meds_section = next_state.get("medications", {}).get("suggestions", [])
                                meds_lines = [m.get("suggestion") for m in meds_section[:5] if m.get("suggestion")]
                                if meds_lines:
                                    summary_lines.append("\n**Treatment suggestions:**\n- " + "\n- ".join(meds_lines))
                                if summary_lines:
                                    append_message("assistant", "\n\n".join(summary_lines))
                            st.rerun()
                        except Exception as e:
                            append_message("assistant", f"Error processing answer: {e}")
                            st.rerun()
            with col2:
                st.caption(q.get("reasoning") or "")

# Results panel (concise design) when diagnosis is complete
def _short_explainer(name: str) -> str:
    mapping = {
        "antacid": "Neutralizes stomach acid for quick relief.",
        "ppi": "Reduces acid production (once daily).",
        "h2": "Reduces acid, often at night.",
        "alginate": "Forms a protective raft to block reflux.",
        "ginger": "May soothe nausea and digestion.",
        "zinc": "Can irritate the esophagus; take with food.",
        "bed elevation": "Reduces nighttime reflux.",
        "weight": "Weight loss reduces abdominal pressure.",
    }
    key = name.lower()
    for k, v in mapping.items():
        if k in key:
            return v
    return "Commonly used option for symptom control."

def render_results(state_dict: dict):
    if not isinstance(state_dict, dict):
        return
    if state_dict.get("status") != "diagnosis_complete":
        return
    final = state_dict.get("final_diagnosis", {})
    meds_section = state_dict.get("medications", {}).get("suggestions", [])

    st.markdown("---")
    st.subheader("🎯 Result")
    c1, c2 = st.columns([2, 1])
    with c1:
        pdx = final.get("primary_diagnosis", "Unknown")
        st.markdown(f"### {pdx}")
        if final.get("final_summary"):
            st.markdown(final.get("final_summary"))
    with c2:
        st.metric("Confidence", f"{int(state_dict.get('confidence_score', 0.0)*100)}%")

    # Medications comparison table (if provided by agent)
    meds_table = state_dict.get("medications", {}).get("options") or state_dict.get("medications", {}).get("suggestions")
    if isinstance(meds_table, list) and meds_table and isinstance(meds_table[0], dict) and "name" in meds_table[0]:
        st.markdown("#### 💊 Medications (ranked)")
        import pandas as pd
        df = pd.DataFrame([
            {
                "Rank": m.get("ranking"),
                "Name": m.get("name"),
                "Family": m.get("family"),
                "Best dose": m.get("best_dose"),
                "OTC/Rx": m.get("otc_or_rx"),
                "Other uses": ", ".join(m.get("other_uses", [])[:4]),
                "One-line": m.get("one_line"),
                "_details": m.get("details"),
            }
            for m in meds_table
        ]).sort_values(by=["Rank", "Name"], na_position="last")
        # Keep details out of main display; use index to show on demand
        show_df = df.drop(columns=["_details"]) if "_details" in df.columns else df
        st.dataframe(show_df, use_container_width=True, hide_index=True)

        # Per-row detail toggles (compact)
        for i, row in df.iterrows():
            if row.get("_details"):
                with st.expander(f"Details: {row['Name']}"):
                    st.write(row["_details"])

    # General suggestions as concise bullets (if present)
    if meds_section and (not (isinstance(meds_table, list) and meds_table and isinstance(meds_table[0], dict) and "name" in meds_table[0])):
        st.markdown("#### 💊 Options")
        for m in meds_section[:8]:
            s = (m or {}).get("suggestion")
            if not s:
                continue
            one_line = _short_explainer(s)
            st.markdown(f"- **{s}** — {one_line}")

    # Export buttons
    st.markdown("---")
    colA, colB = st.columns([1, 1])
    # Build markdown export
    pdx = final.get("primary_diagnosis", "Unknown")
    md_lines = [
        f"# Diagnosis\n\n**Primary:** {pdx}",
        f"**Confidence:** {int(state_dict.get('confidence_score', 0.0)*100)}%",
    ]
    if final.get("final_summary"):
        md_lines.append(f"\n{final.get('final_summary')}")
    if isinstance(meds_table, list) and meds_table and isinstance(meds_table[0], dict) and "name" in meds_table[0]:
        md_lines.append("\n## Medications (ranked)")
        for m in meds_table:
            md_lines.append(f"- {m.get('ranking')}. {m.get('name')} ({m.get('family')}), dose: {m.get('best_dose')} — {m.get('one_line')}")
    elif meds_section:
        md_lines.append("\n## Options")
        for m in meds_section[:8]:
            s = (m or {}).get("suggestion")
            if s:
                md_lines.append(f"- {s}")
    md_content = "\n".join(md_lines)
    with colA:
        st.download_button("Download Markdown", data=md_content.encode("utf-8"), file_name="diagnosis.md", mime="text/markdown")
    with colB:
        if Document is not None:
            try:
                buf = BytesIO()
                doc = Document()
                doc.add_heading("Diagnosis", level=1)
                doc.add_paragraph(f"Primary: {pdx}")
                doc.add_paragraph(f"Confidence: {int(state_dict.get('confidence_score', 0.0)*100)}%")
                if final.get("final_summary"):
                    doc.add_paragraph(final.get("final_summary"))
                if isinstance(meds_table, list) and meds_table and isinstance(meds_table[0], dict) and "name" in meds_table[0]:
                    doc.add_heading("Medications (ranked)", level=2)
                    for m in meds_table:
                        doc.add_paragraph(f"{m.get('ranking')}. {m.get('name')} ({m.get('family')}) — dose: {m.get('best_dose')}\n{m.get('one_line')}")
                        if m.get("details"):
                            doc.add_paragraph(m.get("details"))
                elif meds_section:
                    doc.add_heading("Options", level=2)
                    for m in meds_section[:8]:
                        s = (m or {}).get("suggestion")
                        if s:
                            doc.add_paragraph(s)
                doc.save(buf)
                st.download_button("Download DOCX", data=buf.getvalue(), file_name="diagnosis.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.caption(f"DOCX export unavailable: {e}")

# Auto-render results if available
render_results(st.session_state.get("diag_state"))
